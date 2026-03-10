import pandas as pd
import json
from pathlib import Path
import pytesseract # type: ignore
from PIL import Image # type: ignore
import pdfplumber # type: ignore
import docx
import openpyxl # type: ignore
import xml.etree.ElementTree as ET
import logging

logger = logging.getLogger(__name__)

class UniversalFileParser:

    def parse(self, file_path: str):
        """Parse any supported file format and return extracted text"""
        try:
            ext = Path(file_path).suffix.lower()
            logger.info(f"Parsing file with extension: {ext}")

            if ext == ".pdf":
                return self.parse_pdf(file_path)

            elif ext in [".docx", ".doc"]:
                return self.parse_docx(file_path)

            elif ext == ".csv":
                return self.parse_csv(file_path)

            elif ext in [".xlsx", ".xls"]:
                return self.parse_excel(file_path)

            elif ext == ".txt":
                return self.parse_txt(file_path)

            elif ext == ".json":
                return self.parse_json(file_path)

            elif ext == ".xml":
                return self.parse_xml(file_path)

            elif ext in [".png", ".jpg", ".jpeg"]:
                return self.parse_image(file_path)

            else:
                logger.warning(f"Unsupported file type: {ext}, attempting text parsing")
                return self.parse_txt(file_path)
                
        except Exception as e:
            logger.error(f"Error parsing file {file_path}: {e}")
            raise

    def parse_pdf(self, path):
        try:
            text = ""
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                logger.warning(f"No text extracted from PDF: {path}")
            return text
        except Exception as e:
            logger.error(f"Error parsing PDF {path}: {e}")
            raise

    def parse_docx(self, path):
        try:
            doc = docx.Document(path)
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text += "\n" + row_text
            
            if not text.strip():
                logger.warning(f"No text extracted from DOCX: {path}")
            return text
        except Exception as e:
            logger.error(f"Error parsing DOCX {path}: {e}")
            raise

    def parse_csv(self, path):
        try:
            df = pd.read_csv(path)
            text = f"CSV File with {len(df)} rows and {len(df.columns)} columns\n\n"
            text += "Columns: " + ", ".join(df.columns.astype(str)) + "\n\n"
            
            # Limit to first 1000 rows to avoid huge outputs
            max_rows = min(1000, len(df))
            text += df.head(max_rows).to_string()
            
            if len(df) > max_rows:
                text += f"\n\n... (showing first {max_rows} of {len(df)} rows)"
            
            return text
        except Exception as e:
            logger.error(f"Error parsing CSV {path}: {e}")
            raise

    def parse_excel(self, path):
        try:
            wb = openpyxl.load_workbook(path, data_only=True)
            text = ""
            
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                text += f"\n=== Sheet: {sheet_name} ===\n"
                
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    if row_count >= 1000:  # Limit rows
                        text += "\n... (truncated for brevity)\n"
                        break
                    
                    row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        text += row_text + "\n"
                        row_count += 1
            
            if not text.strip():
                logger.warning(f"No text extracted from Excel: {path}")
            return text
        except Exception as e:
            logger.error(f"Error parsing Excel {path}: {e}")
            raise

    def parse_txt(self, path):
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
            return text
        except Exception as e:
            logger.error(f"Error parsing TXT {path}: {e}")
            raise

    def parse_json(self, path):
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
            return json.dumps(data, indent=2)
        except Exception as e:
            logger.error(f"Error parsing JSON {path}: {e}")
            raise

    def parse_xml(self, path):
        try:
            tree = ET.parse(path)
            root = tree.getroot()
            
            def element_to_text(element, level=0):
                text = "  " * level + f"{element.tag}"
                if element.text and element.text.strip():
                    text += f": {element.text.strip()}"
                text += "\n"
                
                for child in element:
                    text += element_to_text(child, level + 1)
                
                return text
            
            text = f"XML Document: {root.tag}\n\n"
            text += element_to_text(root)
            return text
        except Exception as e:
            logger.error(f"Error parsing XML {path}: {e}")
            raise

    def parse_image(self, path):
        try:
            img = Image.open(path)
            text = pytesseract.image_to_string(img)
            
            if not text.strip():
                logger.warning(f"No text extracted from image: {path}")
                return f"Image file: {Path(path).name} (OCR returned no text)"
            
            return text
        except Exception as e:
            logger.error(f"Error parsing image {path}: {e}")
            raise


parser = UniversalFileParser()