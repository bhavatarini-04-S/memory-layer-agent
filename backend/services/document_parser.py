"""
Document Parser Service for multiple file formats
"""
from pypdf import PdfReader
from docx import Document
import pandas as pd
from typing import Optional
import logging

logger = logging.getLogger(__name__)

class DocumentParser:
    """Parser for different document types"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> Optional[str]:
        """
        Extract text from PDF file
        
        Args:
            file_path: Path to PDF file
            
        Returns:
            Extracted text or None
        """
        try:
            reader = PdfReader(file_path)
            text = ""
            pages_with_text = 0
            total_pages = len(reader.pages)
            
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text and page_text.strip():
                    text += page_text + "\n\n"
                    pages_with_text += 1
            
            # Check if we got any text
            if text.strip():
                logger.info(f"Successfully extracted text from {pages_with_text}/{total_pages} pages")
                return text.strip()
            else:
                # PDF has no extractable text - likely image-based/scanned
                logger.warning(f"PDF has {total_pages} pages but no extractable text - may be image-based or scanned")
                
                # Check if PDF has images
                has_images = False
                for page in reader.pages[:3]:  # Check first 3 pages
                    if '/XObject' in page.get('/Resources', {}):
                        xobject = page['/Resources']['/XObject'].get_object() # type: ignore
                        for obj in xobject:
                            if xobject[obj]['/Subtype'] == '/Image':
                                has_images = True
                                break
                    if has_images:
                        break
                
                if has_images:
                    logger.info("PDF contains images - this is a scanned/image-based PDF")
                    return None  # Will trigger specific error message
                else:
                    logger.info("PDF appears to be empty or uses unsupported encoding")
                    return None
                    
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {e}")
            return None
    
    @staticmethod
    def parse_docx(file_path: str) -> Optional[str]:
        """
        Extract text from DOCX file
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text or None
        """
        try:
            doc = Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        text += row_text + "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {e}")
            return None
    
    @staticmethod
    def parse_txt(file_path: str) -> Optional[str]:
        """
        Extract text from TXT file
        
        Args:
            file_path: Path to TXT file
            
        Returns:
            File contents or None
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {e}")
            return None
    
    @staticmethod
    def parse_csv(file_path: str) -> Optional[str]:
        """
        Extract text from CSV file
        
        Args:
            file_path: Path to CSV file
            
        Returns:
            CSV contents as formatted text or None
        """
        try:
            df = pd.read_csv(file_path)
            
            # Convert DataFrame to readable text
            text = f"CSV File with {len(df)} rows and {len(df.columns)} columns\n\n"
            
            # Add column names
            text += "Columns: " + ", ".join(df.columns) + "\n\n"
            
            # Add row data (limit to prevent huge text)
            max_rows = min(100, len(df))
            for idx, row in df.head(max_rows).iterrows():
                row_text = " | ".join(f"{col}: {val}" for col, val in row.items())
                text += f"Row {idx + 1}: {row_text}\n" # type: ignore
            
            if len(df) > max_rows:
                text += f"\n... and {len(df) - max_rows} more rows"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error parsing CSV {file_path}: {e}")
            return None
    
    @classmethod
    def parse_file(cls, file_path: str, file_type: str) -> Optional[str]:
        """
        Parse file based on its type
        
        Args:
            file_path: Path to file
            file_type: File extension (pdf, docx, txt, csv)
            
        Returns:
            Extracted text or None
        """
        parsers = {
            'pdf': cls.parse_pdf,
            'docx': cls.parse_docx,
            'doc': cls.parse_docx,  # Try docx parser for doc files
            'txt': cls.parse_txt,
            'csv': cls.parse_csv
        }
        
        parser = parsers.get(file_type.lower())
        if not parser:
            logger.error(f"Unsupported file type: {file_type}")
            return None
        
        return parser(file_path)

# Create singleton instance
document_parser = DocumentParser()
